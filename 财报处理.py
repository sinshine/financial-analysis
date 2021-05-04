import pandas as pd
from docx import Document

df1=pd.read_excel('财务报表.xlsx', sheet_name='表1资产负债表')

df1=df1.set_index('类别')

df2=pd.read_excel('财务报表.xlsx', sheet_name='表2利润表')

df2=df2.set_index('类别')

df3=pd.read_excel('财务报表.xlsx', sheet_name='表3现金流量')

df3=df3.set_index('类别')



#资产负债表
df5=df1
df5=df5.fillna(0)
df5=df5.round(2)

#利润表

df6=df2
df6=df6.fillna(0)



#现金流量表


df7=df3
df7=df7.fillna(0)
#df7.loc['期末净现金流量']=df7.loc['经营活动产生的现金流量净额']+df7.loc['投资活动产生的现金流量净额']+df7.loc['筹资活动产生的现金流量净额']+df7.loc['汇率变动影响']


#排序函数
def paixu(list1,df1):
    dict1={}
    for i in list1:
        dict1.update({i:df1.loc[i].sum().round(1)})
    dict1=sorted(dict1.items(),key=lambda item:item[1])
    return dict1

#指标汇总
df0=df5
df0=df0.drop(index=df0.index)
#偿债能力指标
df0.loc['资产负债率']=df5.loc['负债总计']/df5.loc['资产总计']
df0.loc['流动比率']=df5.loc['流动资产合计']/df5.loc['流动负债合计']
df0.loc['速动比率']=(df5.loc['流动资产合计']-df5.loc['存货'])/df5.loc['流动负债合计']
df0.loc['现金比率']=(df5.loc['货币资金']+df5.loc['交易性金融资产'])/df5.loc['流动负债合计']
df0.loc['利息保障倍数']=(df6.loc['四、利润总额（亏损以“－”号填列）']+df6.loc['财务费用'])/df6.loc['财务费用']

#经营周转指标
df0.loc['应收账款周转率（次）']=df6.loc['一、营业总收入']*2/(df5.loc['应收账款']+df5.loc['应收账款'].shift(1))
df0.loc['存货周转率（次）']=df6.loc['其中：营业成本']*2/(df5.loc['存货']+df5.loc['存货'].shift(1))
df0.loc['应付账款周转率（次）']=df6.loc['其中：营业成本']*2/(df5.loc['应付账款']+df5.loc['应付账款'].shift(1))
df0.loc['预收账款周转率（次）']=df6.loc['一、营业总收入']*2/(df5.loc['预收账款']+df5.loc['预收账款'].shift(1))
df0.loc['预付账款周转率（次）']=df6.loc['其中：营业成本']*2/(df5.loc['预付账款']+df5.loc['预付账款'].shift(1))
df0.loc['应收账款周转天数']=360/df0.loc['应收账款周转率（次）']
df0.loc['存货周转天数']=360/df0.loc['存货周转率（次）']
df0.loc['应付账款周转天数']=360/df0.loc['应付账款周转率（次）']
df0.loc['预收账款周转天数']=360/df0.loc['预收账款周转率（次）']
df0.loc['预付账款周转天数']=360/df0.loc['预付账款周转率（次）']
df0.loc['现金周期']=df0.loc['应收账款周转天数']+df0.loc['存货周转天数']-df0.loc['应付账款周转天数']-df0.loc['预收账款周转天数']+df0.loc['预付账款周转天数']
df0.loc['收现率']=df7.loc['销售商品、提供劳务收到的现金']/df6.loc['一、营业总收入']

#盈利指标
df0.loc['成本占比']=df6.loc['其中：营业成本']/df6.loc['一、营业总收入']
df0.loc['毛利润率']=1-df0.loc['成本占比']
df0.loc['营业利润率']=df6.loc['三、营业利润（亏损以“－”号填列）']/df6.loc['一、营业总收入']
df0.loc['销售利润率']=(df6.loc['一、营业总收入']-df6.loc['其中：营业成本']-df6.loc['税金及附加']-df6.loc['销售费用'])/df6.loc['一、营业总收入']
df0.loc['净利润率']=df6.loc['五、净利润（亏损以“－”号填列）']/df6.loc['一、营业总收入']
df0.loc['EBIT']=(df6.loc['五、净利润（亏损以“－”号填列）']+df6.loc['减：所得税费用']+df6.loc['其中：利息费用'])/10000
df0.loc['总资产收益率']=df6.loc['五、净利润（亏损以“－”号填列）']*2/(df5.loc['资产总计']+df5.loc['资产总计'].shift(1))
df0.loc['净资产收益率']=df6.loc['五、净利润（亏损以“－”号填列）']*2/(df5.loc['所有者权益合计']+df5.loc['所有者权益合计'].shift(1))

#报表勾稽关系
df0.loc['报表所得税纳税比例']=df6.loc['减：所得税费用']/df6.loc['四、利润总额（亏损以“－”号填列）']
df0.loc['当年利润与年初年末未分配利润差额的匹配情况']=df5.loc['未分配利润'].shift(1)+df6.loc['五、净利润（亏损以“－”号填列）']-df5.loc['未分配利润']

#财务简表(第一页）
df8=df0
df8=df8.drop(index=df8.index)

#流动资产排序
'''
dict51={'货币资金':df5.loc['货币资金'].sum().round(1),'交易性金融资产':df5.loc['交易性金融资产'].sum().round(1),'应收票据':df5.loc['应收票据'].sum().round(1),'应收账款':df5.loc['应收账款'].sum().round(1),'预付账款':df5.loc['预付账款'].sum().round(1),'应收股利':df5.loc['应收股利'].sum().round(1),'应收利息':df5.loc['应收利息'].sum().round(1),'其他应收款':df5.loc['其他应收款'].sum().round(1),'存货':df5.loc['存货'].sum().round(1),'待摊费用':df5.loc['待摊费用'].sum().round(1),'一年内到期的非流动资产':df5.loc['一年内到期的非流动资产'].sum().round(1),'其他流动资产':df5.loc['其他流动资产'].sum().round(1)}
list51=sorted(dict51.items(),key=lambda item:item[1])
'''
list51=paixu(['交易性金融资产','应收票据','应收账款','预付账款','其他应收款','存货','待摊费用','一年内到期的非流动资产','其他流动资产'],df5)
df8.loc['总资产']=df5.loc['资产总计']
df8.loc['流动资产']=df5.loc['流动资产合计']
df8.loc['货币资金']=df5.loc['货币资金']
df8.loc[list51[-1][0]]=df5.loc[list51[-1][0]]
df8.loc[list51[-2][0]]=df5.loc[list51[-2][0]]
df8.loc[list51[-3][0]]=df5.loc[list51[-3][0]]
df8.loc[list51[-4][0]]=df5.loc[list51[-4][0]]


#非流动资产排序
dict52={'可供出售金融资产':df5.loc['可供出售金融资产'].sum().round(1),'持有至到期投资':df5.loc['持有至到期投资'].sum().round(1),'投资性房地产':df5.loc['投资性房地产'].sum().round(1),'长期股权投资':df5.loc['长期股权投资'].sum().round(1),'长期应收款':df5.loc['长期应收款'].sum().round(1),'固定资产':df5.loc['固定资产'].sum().round(1),'在建工程':df5.loc['在建工程'].sum().round(1),'生产性生物资产':df5.loc['生产性生物资产'].sum().round(1),'油气资产':df5.loc['油气资产'].sum().round(1),'无形资产':df5.loc['无形资产'].sum().round(1),'开发支出':df5.loc['开发支出'].sum().round(1),'商誉':df5.loc['商誉'].sum().round(1),'长期待摊费用':df5.loc['长期待摊费用'].sum().round(1),'递延所得税资产':df5.loc['递延所得税资产'].sum().round(1),'其他非流动资产':df5.loc['其他非流动资产'].sum().round(1)}
list52=sorted(dict52.items(),key=lambda item:item[1])

df8.loc['非流动资产']=df5.loc['非流动资产合计']
df8.loc[list52[-1][0]]=df5.loc[list52[-1][0]]
df8.loc[list52[-2][0]]=df5.loc[list52[-2][0]]
df8.loc[list52[-3][0]]=df5.loc[list52[-3][0]]

#负债排序
dict53={'短期借款':df5.loc['短期借款'].sum().round(1),'交易性金融负债':df5.loc['交易性金融负债'].sum().round(1),'应付票据':df5.loc['应付票据'].sum().round(1),'应付账款':df5.loc['应付账款'].sum().round(1),'预收账款':df5.loc['预收账款'].sum().round(1),'应付职工薪酬':df5.loc['应付职工薪酬'].sum().round(1),'应交税费':df5.loc['应交税费'].sum().round(1),'其他应付款':df5.loc['其他应付款'].sum().round(1),'预计负债':df5.loc['预计负债'].sum().round(1),'一年内到期的长期负债':df5.loc['一年内到期的长期负债'].sum().round(1),'其他流动负债':df5.loc['其他流动负债'].sum().round(1),'长期借款':df5.loc['长期借款'].sum().round(1),'应付债券':df5.loc['应付债券'].sum().round(1),'长期应付款':df5.loc['长期应付款'].sum().round(1),'递延所得税负债':df5.loc['递延所得税负债'].sum().round(1),'其他非流动负债':df5.loc['其他非流动负债'].sum().round(1)}
list53=sorted(dict53.items(),key=lambda item:item[1])
df8.loc['总负债']=df5.loc['负债总计']
df8.loc[list53[-1][0]]=df5.loc[list53[-1][0]]
df8.loc[list53[-2][0]]=df5.loc[list53[-2][0]]
df8.loc[list53[-3][0]]=df5.loc[list53[-3][0]]
df8.loc[list53[-4][0]]=df5.loc[list53[-4][0]]
df8.loc[list53[-5][0]]=df5.loc[list53[-5][0]]

#所有者权益排序
dict54={'实收资本（或股本）':df5.loc['实收资本（或股本）'].sum().round(1),'资本公积':df5.loc['资本公积'].sum().round(1),'盈余公积':df5.loc['盈余公积'].sum().round(1),'一般风险准备':df5.loc['一般风险准备'].sum().round(1),'未分配利润':df5.loc['未分配利润'].sum().round(1)}
list54=sorted(dict54.items(),key=lambda item:item[1])
df8.loc['所有者权益']=df5.loc['所有者权益合计']
df8.loc['资本公积']=df5.loc['资本公积']
df8.loc['未分配利润']=df5.loc['未分配利润']
'''
df8.loc[list54[-1][0]]=df5.loc[list54[-1][0]]
df8.loc[list54[-2][0]]=df5.loc[list54[-2][0]]
df8.loc[list54[-3][0]]=df5.loc[list54[-3][0]]
'''
df8=df8.round(1)

#财务简表(第二页）
df9=df0
df9=df9.drop(index=df9.index)

#偿债能力
df9.loc['资产负债率']=df0.loc['资产负债率'].apply(lambda x:format(x,'.2%'))
df9.loc['流动比率']=df0.loc['流动比率'].round(2)
df9.loc['收现率']=df0.loc['收现率'].round(2)
df9.loc['利息保障倍数']=df0.loc['利息保障倍数'].round(2)

#运营能力
'''
df9.loc['应收账款周转天数']=df0.loc['应收账款周转天数'].round(1)
df9.loc['存货周转天数']=df0.loc['存货周转天数'].round(1)
df9.loc['应付账款周转天数']=df0.loc['应付账款周转天数'].round(1)
df9.loc['预收账款周转天数']=df0.loc['预收账款周转天数'].round(1)
df9.loc['预付账款周转天数']=df0.loc['预付账款周转天数'].round(1)

df9.loc['现金周期']=df0.loc['现金周期'].round(1)
'''
list55=paixu(['应收账款周转天数','存货周转天数','应付账款周转天数','预收账款周转天数','预付账款周转天数'],df0)
df9.loc[list55[-1][0]]=df0.loc[list55[-1][0]].round(1)
df9.loc[list55[-2][0]]=df0.loc[list55[-2][0]].round(1)
df9.loc[list55[-3][0]]=df0.loc[list55[-3][0]].round(1)

#盈利能力
df9.loc['营业收入']=df6.loc['一、营业总收入'].round(1)
df9.loc['营业利润']=df6.loc['三、营业利润（亏损以“－”号填列）'].round(1)
df9.loc['净利润']=df6.loc['五、净利润（亏损以“－”号填列）'].round(1)
df9.loc['EBIT']=df0.loc['EBIT'].round(2)
df9.loc['毛利润率']=df0.loc['毛利润率'].apply(lambda x:format(x,'.2%'))
df9.loc['营业利润率']=df0.loc['营业利润率'].apply(lambda x:format(x,'.2%'))
df9.loc['净利润率']=df0.loc['净利润率'].apply(lambda x:format(x,'.2%'))


#现金流
df9.loc['经营活动现金流入']=df7.loc['经营活动现金流入小计'].round(1)
df9.loc['经营活动现金流出']=df7.loc['经营活动现金流出小计'].round(1)
df9.loc['经营活动净现金流']=df7.loc['经营活动产生的现金流量净额'].round(1)
df9.loc['投资活动净现金流']=df7.loc['投资活动产生的现金流量净额'].round(1)
df9.loc['筹资活动净现金流']=df7.loc['筹资活动产生的现金流量净额'].round(1)
df9.loc['净现金流']=df7.loc['五、现金及现金等价物净增加额'].round(1)


doc2 = Document()

#财务简表
df8=df8[df8.columns[-4:]]
df9=df9[df9.columns[-4:]]
df10=pd.concat([df8.reset_index(),df9.reset_index()],axis=1)
df10.to_excel('财务简表.xlsx',encoding='gbk')

#平均数
def avg(list1):
    avg=sum(list1)/len(list1)
    return avg

#判断
def panduan(df9,text9):
    if df9.loc[text9][-3]<df9.loc[text9][-2]<df9.loc[text9][-1]:
        commit1 = text9+'持续增加，平均增长率为'+'{:.0%}'.format((((df9.loc[text9][-1]/df9.loc[text9][-3])-1)/2))
    elif df9.loc[text9][-3]>df9.loc[text9][-2]>df9.loc[text9][-1]:
        commit1 = text9+'持续下降，平均降幅为'+'{:.0%}'.format((((df9.loc[text9][-1]/df9.loc[text9][-3])-1)/2))
    elif df9.loc[text9][-3]==df9.loc[text9][-2]==df9.loc[text9][-1]==0:
        commit1 = '一直为0。'
    else:
        num = avg([df9.loc[text9][-3],df9.loc[text9][-2],df9.loc[text9][-1]])
        commit1 = '在{:.0%}上下波动。'.format(num)
    return commit1



#偿债能力分析
t1="借款人近三期的总资产分别为%s、%s和%s万元；总负债分别为%s、%s和%s万元；资产负债率分别为%s、%s和%s。整体来看，借款人资产负债率" % (df8.loc['总资产'][-3],df8.loc['总资产'][-2],df8.loc['总资产'][-1],df8.loc['总负债'][-3],df8.loc['总负债'][-2],df8.loc['总负债'][-1],df9.loc['资产负债率'][-3],df9.loc['资产负债率'][-2],df9.loc['资产负债率'][-1]) # 10个变量

if df9.loc['资产负债率'][-3]<df9.loc['资产负债率'][-2]<df9.loc['资产负债率'][-1]:
    commit1 = '持续增加，有加大资本杠杆的趋势。'
elif df9.loc['资产负债率'][-3]>df9.loc['资产负债率'][-2]>df9.loc['资产负债率'][-1]:
    commit1 = '持续下降，资产负债结构有所改善。'
elif df9.loc['资产负债率'][-3]==df9.loc['资产负债率'][-2]==df9.loc['资产负债率'][-1]==0:
    commit1 = '一直为0，近三期均无负债。'
else:
    num = avg([df0.loc['资产负债率'][-3],df0.loc['资产负债率'][-2],df0.loc['资产负债率'][-1]])
    commit1 = '在{:.0%}上下波动。'.format(num)

list56=paixu(['资产总计','货币资金','交易性金融资产','应收票据','应收账款','预付账款','其他应收款','存货','待摊费用','一年内到期的非流动资产','其他流动资产','可供出售金融资产', '持有至到期投资', '投资性房地产', '长期股权投资', '长期应收款', '固定资产', '在建工程',  '生产性生物资产', '油气资产', '无形资产', '开发支出', '商誉', '长期待摊费用', '递延所得税资产', '其他非流动资产'],df5)


t5="借款人近三期流动资产分别为%s、%s和%s万元，在总资产构成中流动资产占比分别%.2f%%、%.2f%%和%.2f%%。主要资产构成为%s、%s、%s、%s、%s，在总资产构成中的占比分别为%.2f%%、%.2f%%、%.2f%%、%.2f%%、%.2f%%。【具体详见重点科目分析】"% (df8.loc['流动资产'][-3],df8.loc['流动资产'][-2],df8.loc['流动资产'][-1],df8.loc['流动资产'][-3]/df8.loc['总资产'][-3]*100,df8.loc['流动资产'][-2]/df8.loc['总资产'][-2]*100,df8.loc['流动资产'][-1]/df8.loc['总资产'][-1]*100,list56[-2][0],list56[-3][0],list56[-4][0],list56[-5][0],list56[-6][0],list56[-2][1]/list56[-1][1]*100,list56[-3][1]/list56[-1][1]*100,list56[-4][1]/list56[-1][1]*100,list56[-5][1]/list56[-1][1]*100,list56[-6][1]/list56[-1][1]*100,)

t2="从债务期限结构看，借款人近三期流动负债分别为%s、%s和%s万元，流动负债占比分别为%.2f%%、%.2f%%和%.2f%%。总体来看，公司偿债能力【较大\在合理范围内\较小，根据实际情况自行评价】。"  % (df5.loc['流动负债合计'][-3],df5.loc['流动负债合计'][-2],df5.loc['流动负债合计'][-1],df5.loc['流动负债合计'][-3]/df5.loc['负债总计'][-3]*100,df5.loc['流动负债合计'][-2]/df5.loc['负债总计'][-2]*100,df5.loc['流动负债合计'][-1]/df5.loc['负债总计'][-1]*100,)

#盈利能力分析
t3="借款人前三期分别实现营业收入%s、%s和%s万元，%s。毛利率分别为%.2f%%、%.2f%%和%.2f%%，实现净利润%s、%s和%s万元，净利润率分别为%.2f%%、%.2f%%和%.2f%%，总体来看，借款人盈利能力【相对较好/稳定/较弱，根据实际情况自行评价】"  % (df9.loc['营业收入'][-3],df9.loc['营业收入'][-2],df9.loc['营业收入'][-1],panduan(df9,'营业收入'),df0.loc['毛利润率'][-3]*100,df0.loc['毛利润率'][-2]*100,df0.loc['毛利润率'][-1]*100,df9.loc['净利润'][-3],df9.loc['净利润'][-2],df9.loc['净利润'][-1],df0.loc['净利润率'][-3]*100,df0.loc['净利润率'][-2]*100,df0.loc['净利润率'][-1]*100)


#营运能力分析
t6="借款人近三期现金周期分别为%s、%s和%s。其中最新一期%s为%s，%s为%s，%s为%s。近三期流动比率分别为%s、%s和%s，三期平均值为%s；借款人近三期速动比率分别为%s、%s和%s，三期平均值为%s。【请根据实际情况自行评价】。"  % (df0.loc['现金周期'][-3].round(1),df0.loc['现金周期'][-2].round(1),df0.loc['现金周期'][-1].round(1),list55[-1][0],df9.loc[list55[-1][0]][-1],list55[-2][0],df9.loc[list55[-2][0]][-1],list55[-3][0],df9.loc[list55[-3][0]][-1],df0.loc['流动比率'][-3].round(1),df0.loc['流动比率'][-2].round(1),df0.loc['流动比率'][-1].round(1),avg([df0.loc['流动比率'][-3],df0.loc['流动比率'][-2],df0.loc['流动比率'][-1]]).round(1),df0.loc['速动比率'][-3].round(1),df0.loc['速动比率'][-2].round(1),df0.loc['速动比率'][-1].round(1),avg([df0.loc['速动比率'][-3],df0.loc['速动比率'][-2],df0.loc['速动比率'][-1]]).round(1))


#现金流量情况
t4="借款人近三期分别实现净现金流入%s、%s和%s万元。其中：经营活动现金净流入%s、%s和%s万元；投资活动现金净流入%s、%s和%s万元；筹资活动现金净流入%s、%s和%s万元。总体来看，借款人现金流情况【请根据实际情况自行评价】" % (df7.loc['五、现金及现金等价物净增加额'][-3].round(1),df7.loc['五、现金及现金等价物净增加额'][-2].round(1),df7.loc['五、现金及现金等价物净增加额'][-1].round(1),df9.loc['经营活动净现金流'][-3],df9.loc['经营活动净现金流'][-2],df9.loc['经营活动净现金流'][-1],df9.loc['投资活动净现金流'][-3],df9.loc['投资活动净现金流'][-2],df9.loc['投资活动净现金流'][-1],df9.loc['筹资活动净现金流'][-3],df9.loc['筹资活动净现金流'][-2],df9.loc['筹资活动净现金流'][-1],)

#小结
t7="整体来看，借款人偿债能力【】，盈利能力【】，运营能力【】，现金流情况【】，整体财务状况【】。"



doc2.add_heading('一、财务报表分析', 0)
doc2.add_paragraph('关注并表范围是否有变化！对企业影响比较大的会计政策和会计估计是否有变化！')
doc2.add_heading('偿债能力分析', 0)
doc2.add_paragraph('重点关注资产、负债、权益的结构，以及是否存在需要还原！')
t1=doc2.add_paragraph(t1)
t1.add_run(commit1)
t1.add_run(t5)
t1.add_run(t2)

doc2.add_paragraph('变现性：对于流动资产尤为重要。应关注应收账款的账龄，判断其回收的难度。存货的周转及变现能力等。')
doc2.add_paragraph('被利用性：主要针对非流动资产。固定资产应关注其产能利用率，如果长期闲置则无法产生收入，拖累企业运营。无形资产应关注其构成，一般土地、特许经营权都是比较有价值的，专利等则要具体分析。')
doc2.add_paragraph('增值性：资产主要采用历史成本法计价，要关注资产账面价值与实际价值之间的差异。土地可能比账面价值更值钱，而机器设备、专用厂房等由于专用性较强，可变现价值往往远低于账面价值。')
doc2.add_paragraph('与其他资产组合的增值能力：企业的资产之间应有相互关联，部分企业多元化投资倾向较强，产业之间没有协同效应，风险程度比较高。')
doc2.add_paragraph('与负债的期限匹配情况：资产和负债的期限应基本匹配，流动负债主要用于流动资产，若短债长用情况太严重，可能产生流动性风险。')



doc2.add_heading('盈利能力分析', 0)
doc2.add_paragraph(t3)
doc2.add_paragraph('关注利润结构，依靠自身经营获取的利润是否是主流（投资收益-靠别人经营，补贴收入-靠政府，营业外收入-靠运气）。关注利润率变化情况，与同业对比，过高或过低的利润率都应有合理解释。')
doc2.add_paragraph('利息费用：利息费用/有息负债余额可以大致得到企业融资成本。毛利/经营活动净现金流高于利息费用是底线，否则企业需要借款还利息，类似庞氏骗局。在建工程产生的利息费用资本化，没有体现在利润表里，只能从现金流量表中找。')

doc2.add_heading('营运能力分析', 0)
doc2.add_paragraph(t6)
doc2.add_heading('现金流情况', 0)
doc2.add_paragraph(t4)
doc2.add_paragraph('经营活动现金流：是收付实现制下的净利润，是归还固定资产贷款、融资利息、扩大再生产、归还经营性负债、分红的来源。注意排除其他应收款的扰动-收到、支付其他与经营活动有关的现金科目，剔除关联方相关资金扰动后再看造现能力。')
doc2.add_paragraph('投资活动现金流：体现企业的扩张和收缩策略。对内扩张-固定资产、无形资产。对外扩张-股权投资。投资时要承担风险的-增加未来的不确定性。投资是要本钱的-挤占自有资金，带来新的负债。')
doc2.add_paragraph('筹资活动现金流：吸收权益性资金-降低杠杆。对外借款-提高杠杆。结合取得借款收到、发行债券收到、偿还债务支付三个科目看债务资金是净流入还是净流出。分配股利、利润或偿付利息支付的科目把股权资金和债权资金的回报混在一起，带来了分析的障碍，应结合资产负债表中未分配利润和利润表中净利润的值进行判断。')

doc2.add_heading('小结', 0)
doc2.add_paragraph('财务报表中异常情况如果没有合理解释，很可能是财务粉饰造假的信号。')
doc2.add_paragraph('可利用其他信息对财务报表进行验证，如用水、用电、职工工资、税收等外部信息，如银行流水判断经营情况，如通过上下游客户侧面了解其情况。')
doc2.add_paragraph(t7)
doc2.add_heading('二、重点科目分析', 0)
doc2.add_paragraph('货币资金：关注受限制货币资金，受限资金一般是债务融资保证金，也可能是资本市场募集的有确定用途的股本金。在考虑企业流动性的时候应该将其剔除。关注存贷双高，拥有大量货币资金的同时还存在大量借款（做套利-存贷利差，内保外贷。钱刚刚借来还没用。）')
doc2.add_paragraph('应收账款：关注应收账款的账龄，判断其回收的难度。账龄结构-看减值准备计提是否合理。对手方情况-若集中度较高，应分析对手方资信情况。增速是否与收入增速匹配-如增速高于收入增速，说明企业通过延长账期来获取收入，供应链地位一般，甚至有财务粉饰造假可能。')
doc2.add_paragraph('存货：存货分为原材料、在产品、产成品、商品等类型，应分析存货结构，存货余额增幅等，关注存货的周转及变现能力。如果存货余额大增，存货周转率降低，要了解企业是因为销售不畅导致存货积压，还是在策略性的囤货（由于原材料存在上涨趋势而提前备货）？（建筑行业的工程施工已完成未结算在新报表已调整到合同资产）')
doc2.add_paragraph('其他应收款：和负债的其他应付款一样，重点关注与母公司、关联方的往来，一般为大股东占款。通过合并、单一报表的数据对比，看出占款是给并表范围内企业还是并表范围外企业。新报表包含应收利息、应收股利。')
doc2.add_paragraph('固定资产：固定资产应关注其产能利用率，如果长期闲置则无法产生收入，拖累企业运营。关注固定资产的规模、先进程度，若原值大、净值低，说明设备老化，未来可能要发生资产更新支持。审核固定资产折旧方法、折旧年限是否合理，是否遵守一贯性原则。')
doc2.add_paragraph('在建工程：关注项目是否真实存在-在建工程经常作为造假的工具。规模如何-在建工程是新建投资，未来能否完工、达产，具有不确定性。融资结构如何-未来是否会造成债务压力。')
doc2.add_paragraph('长期股权投资：关注规模-生产型企业的长期股权投资不应过大。投向-投资标的是否能与企业本身产生协同效益，过度多元化风险高。与负债的匹配-如长期股权投资高于所有者权益，说明有一部分债务资金被投入了股权投资。')
doc2.add_paragraph('无形资产：关注规模-无形资产占比不应过大。内容-土地所有权、特许经营权通常价值比较高，可以带来比较确定的现金流入，商标权、专利权等则应具体分析。')
doc2.add_paragraph('商誉：并购中收购价款高于标的所有权权益的部分，只存在于合并报表中，没有并购的企业是没有商誉的。商誉不摊销，但需要做减值测试，如果商誉占比比较高，应研究并购标的的经营情况是否如预期一样好，或是否能为企业带来协同效益。如果为否，应考虑将其还原。')
doc2.add_paragraph('预付账款：体现公司在产业链中的行业地位、业务模式和话语权，同时与存货及商品价格有关。关注预付账款应为核心业务相关的，巨额预付应关注。')


doc2.save('财务报表分析.docx')


